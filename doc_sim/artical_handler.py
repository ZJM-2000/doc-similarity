'''
将docx.pdf均转换为txt
@:param  artical_directory
by: Junyi
'''
from docx import Document
import os
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfpage import PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.layout import *
from pdfminer.converter import PDFPageAggregator


class ArticalHandler(object):

    def __init__(self, artical_directory=None):
        self.artical_directory = artical_directory

    @staticmethod
    def change_docx_to_txt(docx_r_path,docx_w_path):
        outputs=open(docx_w_path,'w',encoding='UTF-8')
        document=Document(docx_r_path)
        for paragraph in document.paragraphs:
            if paragraph.text:
                outputs.write(paragraph.text)

    @staticmethod
    def change_pdf_to_txt(pdf_r_path,pdf_w_path):
        fp = open(pdf_r_path, "rb")
        # 来创建一个pdf文档分析器
        parser = PDFParser(fp)
        # 创建一个PDF文档对象存储文档结构
        document = PDFDocument(parser)
        # 检查文件是否允许文本提取
        if not document.is_extractable:
            raise PDFTextExtractionNotAllowed
        else:
            # 创建一个PDF资源管理器对象来存储共赏资源
            rsrcmgr = PDFResourceManager()
            # 设定参数进行分析
            laparams = LAParams()
            # 创建一个PDF设备对象
            # device=PDFDevice(rsrcmgr)
            device = PDFPageAggregator(rsrcmgr, laparams=laparams)
            # 创建一个PDF解释器对象
            interpreter = PDFPageInterpreter(rsrcmgr, device)

            f = open(pdf_w_path, 'w', encoding='UTF-8')
            # 处理每一页
            for page in PDFPage.create_pages(document):
                # print(page)
                interpreter.process_page(page)
                # 接受该页面的LTPage对象
                layout = device.get_result()
                for x in layout:
                    if (isinstance(x, LTTextBoxHorizontal)):
                        txt = x.get_text()
                        #print(txt)
                        f.write(txt)
